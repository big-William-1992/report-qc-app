# -*- coding: utf-8 -*-
"""
生成《星衍AI放射质控·使用说明书》Word 文档。
运行：python3 tools/gen_manual_docx.py
输出：使用说明书.docx（项目根目录）
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "使用说明书.docx")

# 配色
BLUE = "1B3F7A"     # 主蓝
BLUE_L = "E7EFFD"   # 浅蓝底
GREEN = "0E8C7E"
RED = "B71C1C"
ORANGE = "8A5A00"
GRAY = "5A6E8C"
BG_LIGHT = "F2F6FD"
SEV_HIGH = "FDE8E8"
SEV_MED = "FEF3DD"
SEV_LOW = "EDF1F6"


def set_run(run, text=None, size=10.5, bold=False, color=None, font="微软雅黑"):
    if text is not None:
        run.text = text
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=6, space_before=0, font="微软雅黑"):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(), text, size=size, bold=bold, color=color, font=font)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor.from_string(BLUE)
    return h


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell(cell, text, size=9.5, bold=False, color=None, align="left", fill=None, font="微软雅黑"):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    set_run(p.add_run(), text, size=size, bold=bold, color=color, font=font)
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows, col_widths=None, header_fill=BLUE, header_color="FFFFFF",
              zebra=True, font_size=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, size=font_size, bold=True, color=header_color,
                 align="center", fill=header_fill)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            fill = BG_LIGHT if (zebra and i % 2 == 1) else None
            set_cell(t.rows[i + 1].cells[j], val, size=font_size, fill=fill)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


def add_step(doc, num, title, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    set_run(p.add_run(), "Step %d  " % num, size=11, bold=True, color=BLUE)
    set_run(p.add_run(), title, size=11, bold=True)
    if desc:
        d = doc.add_paragraph()
        d.paragraph_format.space_after = Pt(6)
        d.paragraph_format.left_indent = Cm(0.6)
        set_run(d.add_run(), desc, size=10, color=GRAY)


def add_mock_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(), "▎界面示意：" + text, size=9, bold=True, color=GRAY)


# ============================================================
def main():
    doc = Document()

    # 全局默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()
    add_para(doc, "星衍AI放射质控系统", size=28, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "使用说明书", size=22, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_para(doc, "—— 放射科报告智能质控助手 ——", size=12, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_para(doc, "面向放射科医生：粘贴即查 · 复制即控 · 离线OCR自动回填 · 红/橙/灰高亮 · 样本沉淀 · 科室报表",
             size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "所有数据仅存本机，不上传任何服务器", size=10, color=GREEN,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------- 目录 ----------
    add_heading(doc, "目录", 1)
    toc_items = [
        ("一", "软件简介"),
        ("二", "快速上手（三步完成一次质控）"),
        ("三", "报告质控页详解"),
        ("四", "智能采集：剪贴板监听与屏幕 OCR"),
        ("五", "质控看板与统计报表"),
        ("六", "样本库"),
        ("七", "RIS 直连（院内可选）"),
        ("八", "规则维护"),
        ("九", "系统设置与快捷键"),
        ("十", "数据安全与合规"),
        ("十一", "常见问题（FAQ）"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(), "%s、%s" % (num, title), size=11, color=BLUE)
    doc.add_page_break()

    # ---------- 一、软件简介 ----------
    add_heading(doc, "一、软件简介", 1)
    add_para(doc, "星衍AI放射质控是一款面向放射科医生的报告智能质控工具。您只需把影像报告粘贴进来、或从 PACS 屏幕框选抓取，"
                  "系统便会自动完成质控：定位错别字、左右混淆、描述-结论矛盾、评分缺失、部位不符等问题，并按严重度红/橙/灰高亮，"
                  "同时给出多维评分，帮助把好每一份报告的质量关。", size=11)
    add_para(doc, "三大核心能力：", size=11, bold=True)
    add_table(doc,
              ["能力", "说明"],
              [
                  ["粘贴即质控", "复制/粘贴报告全文，自动分栏描述与诊断，立即出质控结果"],
                  ["屏幕识别", "框选 PACS 患者信息栏/报告区域，本地 OCR 离线识别并自动回填"],
                  ["19 项规则", "错别字、左右混淆、描述结论矛盾、缺失随访建议、部位不符等自动检测"],
              ],
              col_widths=[3.5, 12.0])

    # ---------- 二、快速上手 ----------
    add_heading(doc, "二、快速上手（三步完成一次质控）", 1)
    add_step(doc, 1, "启动并登录",
             "双击桌面快捷方式（Windows 绿色版 exe / macOS 启动器）启动软件。首次使用需依次完成：免责声明 → 创建首个账号（工号+密码）"
             "→ 登录。试用期结束后需在「设置 → 授权与激活」输入激活码。")
    add_step(doc, 2, "录入报告",
             "在「报告质控」页填写或自动识别元信息（患者/性别/年龄/成像方式/检查部位/侧别），再输入或粘贴「影像描述」与「影像结论」；"
             "也可直接点「📋 粘贴全文并质控」，一次粘贴报告全文自动分栏。")
    add_step(doc, 3, "运行质控并查看结果",
             "点「▶ 运行质控」（快捷键 Ctrl+Enter）。右侧「质控发现」按严重度列出问题，正文中对应位置高亮，可点「✨ 一键采纳修正」"
             "自动修正确定性的错别字；确认后点「💾 入库」沉淀到样本库。")

    add_mock_title(doc, "报告质控页布局")
    # 界面示意图：用表格构建报告质控页
    add_para(doc, "（左侧为报告录入区，右侧为质控结果区）", size=8.5, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    t = doc.add_table(rows=3, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 第一行：元信息
    m = t.rows[0].cells[0].merge(t.rows[0].cells[1])
    set_cell(m, "元信息卡：姓名 / 性别 / 年龄 / 成像方式 / 检查部位 / 侧别", bold=True, fill=BLUE_L)
    # 第二行：左输入 / 右结果
    set_cell(t.rows[1].cells[0], "影像描述（Findings）文本区", fill=BG_LIGHT)
    set_cell(t.rows[1].cells[1], "多维评分（准确性/完整性/规范性/及时性）", fill=SEV_LOW)
    set_cell(t.rows[2].cells[0], "影像结论（Impression）文本区\n\n操作条：▶运行质控 ✨采纳修正 📄导出报告单 💾入库 📷框选设置 ⚡一键识别 📥加入队列",
             fill=BG_LIGHT)
    set_cell(t.rows[2].cells[1], "质控发现列表\n⛔严重 1 · ⚠警告 1 · ℹ提示 2\n（问题列表 / 原文标注 两个页签）",
             fill=SEV_LOW)
    add_para(doc, "", size=4, space_after=2)

    # ---------- 三、报告质控页详解 ----------
    add_heading(doc, "三、报告质控页详解", 1)
    add_para(doc, "这是您每天最常用的页面。界面分为「录入区（左）」与「结果区（右）」。", size=11)
    add_heading(doc, "3.1 元信息栏", 2)
    add_para(doc, "填写患者姓名、性别、年龄、成像方式、检查部位、侧别。可手工填写，也可通过 OCR 或自动识别回填。"
                  "系统会用它来核验「信息框-正文」一致性（如性别与前列腺/子宫等器官是否冲突）。", size=10.5)
    add_heading(doc, "3.2 影像描述 与 影像结论", 2)
    add_para(doc, "两栏分别对应报告的「检查所见（Findings）」与「诊断印象（Impression）」。系统按栏位进行描述↔结论的交叉比对，"
                  "捕获「描述正常、结论却写病灶」或「左右侧不一致」等隐蔽逻辑错误。", size=10.5)
    add_heading(doc, "3.3 运行质控与结果", 2)
    add_para(doc, "点击「运行质控」后，右侧展示：", size=10.5)
    add_table(doc,
              ["区域", "作用"],
              [
                  ["多维评分", "准确性 / 完整性 / 规范性 / 及时性 四维评分，并附逐项扣分依据，透明可复核"],
                  ["质控发现", "按严重度排序列出每条问题：规则号 + 类型 + 严重度 + 说明 + 原文片段"],
                  ["原文标注", "切换到「原文标注」页签，可直接在描述/诊断原文中看到高亮位置"],
              ],
              col_widths=[4.0, 11.5])
    add_para(doc, "严重度颜色对照：", size=10.5, bold=True)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    set_cell(t.rows[0].cells[0], "严重度", align="center", bold=True, fill=BLUE, color="FFFFFF")
    set_cell(t.rows[0].cells[1], "颜色", align="center", bold=True, fill=BLUE, color="FFFFFF")
    set_cell(t.rows[0].cells[2], "示例规则", align="center", bold=True, fill=BLUE, color="FFFFFF")
    set_cell(t.rows[1].cells[0], "严重 ⛔", fill=SEV_HIGH)
    set_cell(t.rows[1].cells[1], "红", align="center", fill=SEV_HIGH)
    set_cell(t.rows[1].cells[2], "左右混淆、描述-结论矛盾、性别矛盾", fill=SEV_HIGH)
    set_cell(t.rows[2].cells[0], "警告 ⚠", fill=SEV_MED)
    set_cell(t.rows[2].cells[1], "橙", align="center", fill=SEV_MED)
    set_cell(t.rows[2].cells[2], "同音错别字、评分缺失、部位不符", fill=SEV_MED)
    set_cell(t.rows[3].cells[0], "提示 ℹ", fill=SEV_LOW)
    set_cell(t.rows[3].cells[1], "灰", align="center", fill=SEV_LOW)
    set_cell(t.rows[3].cells[2], "单位错误、模板合规、随访建议", fill=SEV_LOW)

    # ---------- 四、智能采集 ----------
    add_heading(doc, "四、智能采集：剪贴板监听与屏幕 OCR", 1)
    add_heading(doc, "4.1 剪贴板监听（复制即质控）", 2)
    add_para(doc, "在「⚙ 高级设置」中开启「监听剪贴板」。之后每当您复制报告（≥15 字），系统自动分栏填入描述/结论并立即质控，"
                  "命中问题会弹窗提醒，听写误写及时拦截。", size=10.5)
    add_heading(doc, "4.2 屏幕 OCR（框选 PACS）", 2)
    add_para(doc, "点击「📷 框选设置」或按 Ctrl+Shift+O 打开框选窗口：", size=10.5)
    add_step(doc, 1, "把 PACS 报告窗口放到前台，点「🖥 截取 PACS 画面」。", None)
    add_step(doc, 2, "用三个彩色框分别框选：病人基础信息（蓝）/ 影像描述（绿）/ 影像诊断（橙）。", None)
    add_step(doc, 3, "点「💾 记住框位」保存位置，下次自动复原；点「⚡ 识别·导入·质控」一键完成。", None)
    add_para(doc, "OCR 为本地离线识别（RapidOCR），模型内置，不联网、截图不出本机，全程数据不出域。", size=10, color=GREEN)
    add_heading(doc, "4.3 后台快捷键（一键质控）", 2)
    add_para(doc, "在「设置 → 快捷键」中可配置并重绑快捷键。配置后，即使焦点在 PACS 等其它窗口，按下快捷键也能触发一键质控。", size=10.5)

    # ---------- 五、质控看板 ----------
    add_heading(doc, "五、质控看板与统计报表", 1)
    add_para(doc, "点击左侧「📊 质控看板」，汇总已入库样本的质控数据：", size=11)
    add_table(doc,
              ["模块", "内容"],
              [
                  ["统计卡", "累计质检报告、今日新增、本周新增、通过率"],
                  ["图表区", "成像方式分布、错误类型分布、质控量趋势（近 30 天）"],
                  ["统计报表", "问题类型 TOP 榜、规则命中 TOP 榜、医生排行榜（按问题数），可切换近 7 天/30 天/季度/全部"],
              ],
              col_widths=[4.0, 11.5])
    add_mock_title(doc, "质控看板顶部统计卡")
    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    cards = [("累计质检报告", "1,286", "↑ 本周 +32", BLUE_L),
             ("今日新增", "12", "↑ 较昨日 +3", "E8F7F3"),
             ("本周新增", "87", "↑ 周环比 +15%", "FDF3E7"),
             ("通过率", "92%", "↓ 昨日 -2%", "FDE8E8")]
    for j, (label, val, trend, fill) in enumerate(cards):
        c = t.rows[0].cells[j]
        set_cell(c, label, align="center", size=9, fill=fill)
    for j, (label, val, trend, fill) in enumerate(cards):
        set_cell(t.rows[1].cells[j], val + "\n" + trend, align="center", bold=True, size=11, fill=fill)

    # ---------- 六、样本库 ----------
    add_heading(doc, "六、样本库", 1)
    add_para(doc, "「📦 样本库」沉淀所有质控过的报告。可管理样本、查看详情、导出/导入样本库数据；"
                  "支持导出为 CSV / JSON 批量数据，或对单份报告导出 Word/PDF 质控报告单。"
                  "建议开启「入库时脱敏」，去除患者姓名等隐私信息后再入库。", size=11)

    # ---------- 七、RIS 直连 ----------
    add_heading(doc, "七、RIS 直连（院内可选）", 1)
    add_para(doc, "若院内网络允许，可在「🔗 RIS 直连」页配置数据库连接并批量拉取报告质控入库。该功能需由院内 IT 提供连接参数与查询 SQL，"
                  "配置仅保存在本机，仅在院内内网使用。", size=11)
    add_step(doc, 1, "填写数据库类型（SQL Server / Oracle / MySQL / PostgreSQL）、主机、端口、库名、账号密码。", None)
    add_step(doc, 2, "粘贴 IT 提供的查询 SQL（须返回 report_text 字段），点「🔌 测试连接」。", None)
    add_step(doc, 3, "点「📥 拉取报告」，可将结果「发送到质控」「批量质控入库」或「全部加入队列」。", None)
    add_para(doc, "还可开启「⏰ 主动轮询质检」：设置拉取间隔与上限，系统定时自动拉取新报告、质控入库，实现「发现即质控」。", size=10.5)

    # ---------- 八、规则维护 ----------
    add_heading(doc, "八、规则维护", 1)
    add_para(doc, "「⚙️ 规则维护」页可查看所有质控规则（规则ID/名称/分类/严重度/状态），并维护词表，保存后立即生效：", size=11)
    add_table(doc,
              ["维护项", "说明"],
              [
                  ["错别字词库（R8）", "维护「错词=正词」词条，支持搜索/新增/批量导入/停用/删除"],
                  ["逻辑错误（R9）", "维护「词A|词B」互斥冲突，可设范围（正文/同一句/描述段/结论段/描述vs结论）"],
                  ["忽略词白名单", "维护命中即不报的词/短语，减少误报"],
                  ["模板规范（R10）", "开关：是否要求给出随访/复查建议"],
                  ["错字检测（R19）", "开启/关闭，并可调灵敏度：低=仅同音 / 中=同音+近音 / 高=近音+形近"],
                  ["智能学习", "扫描历史报告词频，自动发现「低频写法→高频标准写法」候选错字，一键采纳"],
              ],
              col_widths=[5.0, 10.5])
    add_para(doc, "提示：规则词表写入 rules_config.json，桌面端与 Web 端共用；「恢复默认」可还原出厂规则库。", size=9.5, color=GRAY)

    # ---------- 九、设置 ----------
    add_heading(doc, "九、系统设置与快捷键", 1)
    add_para(doc, "点击顶栏「⚙️」打开系统设置：", size=11)
    add_table(doc,
              ["分组", "可配置项"],
              [
                  ["基础", "操作工号（责任到人）、默认成像方式、OCR 置信度阈值、界面主题（浅/深色）"],
                  ["自动化", "OCR 回填后自动质控、屏幕采集/RIS 拉取自动入队、识别前重新抓屏、动态识别、静默质控、剪贴板监听"],
                  ["隐私", "入库时脱敏（姓名等 PHI 不落库）"],
                  ["快捷键", "运行质控 / 存样本库 / 框选OCR / 主题切换，均可点击「重绑」自定义"],
                  ["授权", "查看授权状态、机器识别码（发卡用）、输入激活码激活"],
              ],
              col_widths=[3.5, 12.0])
    add_heading(doc, "快捷键速查表", 2)
    add_table(doc,
              ["功能", "默认快捷键"],
              [
                  ["运行质控", "Ctrl + Enter"],
                  ["存入样本库", "Ctrl + S"],
                  ["框选 OCR · 一键识别", "Ctrl + Shift + O"],
                  ["明暗主题切换", "Ctrl + T"],
                  ["清空录入", "Esc"],
              ],
              col_widths=[8.0, 7.0])

    # ---------- 十、数据安全 ----------
    add_heading(doc, "十、数据安全与合规", 1)
    add_para(doc, "本软件为报告质量辅助核查工具，不替代医师诊断。涉及二类医疗器械与等保三级相关合规要求，请在使用前完成相应注册与测评。", size=11)
    add_table(doc,
              ["项目", "说明"],
              [
                  ["数据存储", "所有报告、样本、配置、账号仅存本机 SQLite（开发态 assets/，打包态 %APPDATA%/MedicalReportQC/）"],
                  ["网络", "不发起任何网络上传；OCR 为本地推理，截图不出本机"],
                  ["脱敏", "建议在共享/公用工作站启用「入库时脱敏」，并定期清理样本库"],
                  ["RIS 直连", "仅连接院内内网数据库，凭据存本机，不上云"],
              ],
              col_widths=[3.5, 12.0])

    # ---------- 十一、FAQ ----------
    add_heading(doc, "十一、常见问题（FAQ）", 1)
    add_table(doc,
              ["问题", "处理办法"],
              [
                  ["复制了报告没触发质控？", "确认已在设置开启「监听剪贴板」；复制内容需 ≥15 字；macOS 需允许访问剪贴板"],
                  ["OCR 识别不准？", "框选区域需包含清晰文字；识别前可勾选「重新抓屏」；调高设置里的 OCR 置信度阈值更严格"],
                  ["自动修正改错了？", "自动修正只改确定性的同音错别字且先预览；矛盾类绝不自动改写。误报可在忽略词白名单中固化"],
                  ["导出报表中文乱码？", "报表为 UTF-8-BOM，用 Excel 直接打开正常"],
                  ["怎么批量处理已有报告？", "用 RIS 直连批量拉取，或用屏幕 OCR 抓取 PACS 报告，或逐份粘贴/导入"],
                  ["忘记账号密码？", "账号为本地库；忘记管理员密码可重置本地 accounts.db，重置后首次启动重新创建管理员"],
              ],
              col_widths=[5.0, 10.5])

    add_para(doc, "", space_after=6)
    add_para(doc, "—— 说明书完 · 如遇问题请联系系统管理员 ——", size=9, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUT)
    print("已生成：", os.path.abspath(OUT))


if __name__ == "__main__":
    main()
